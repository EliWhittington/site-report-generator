import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageOps
from datetime import date
import os
import io
import re
import tempfile

# ==== Streamlit App Config ====
st.set_page_config(page_title="Site Report Generator", layout="wide")
st.title("📷 Site Observation Report Generator")

# === Helper: Extract number from filename for sorting ===
def extract_number(filename):
    match = re.search(r'(\d+)', filename)
    return int(match.group(1)) if match else float('inf')

# === Resize image ===
def resize_image(image_bytes, max_dim):
    img = Image.open(io.BytesIO(image_bytes))
    img = ImageOps.exif_transpose(img)
    img = img.convert("RGB")
    img.thumbnail((max_dim, max_dim))
    return img

# === Add footer with page number ===
def add_footer_with_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.extend([fldChar1, instrText, fldChar2, fldChar3])

# === Generate the report ===
def generate_report(image_paths, weather, subcontractors, areas):
    doc = Document()

    # Add footer with page numbers
    section = doc.sections[0]
    add_footer_with_page_number(section)

    # === Title Page ===
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = 1  # Center
    run = title_paragraph.add_run("Project Observation Report")
    run.font.name = 'Calibri'
    run.font.size = Pt(28)
    doc.add_paragraph()

    today_str = date.today().strftime("%B %d, %Y")

    # === Static Fields ===
    project_name = "National Fire Cache Building                                        -"
    user_name = "Eli Whittington"
    email = "eli.whittington@pc.gc.ca"
    address = "200 Hawk Ave, Banff AB T1L1K2"

    def add_field(label, value):
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        run_label.font.name = 'Calibri'
        run_label.font.size = Pt(12)
        run_value = p.add_run(value)
        run_value.font.name = 'Calibri'
        run_value.font.size = Pt(12)

    add_field("Project", project_name)
    add_field("Name", user_name)
    add_field("Email", email)
    add_field("Review Date", today_str)
    add_field("Report Date", today_str)
    add_field("Address", address)
    add_field("Conditions", f"{weather} °C")

    doc.add_paragraph()  # spacing

    doc.add_paragraph("Subcontractors Present:", style="Heading 2")
    for s in subcontractors:
        doc.add_paragraph(f"    - {s}")

    doc.add_paragraph("Areas of Work/Progress:", style="Heading 2")
    for a in areas:
        doc.add_paragraph(f"    - {a}")

    doc.add_page_break()

    # === Insert 2 Images Per Page (no blank pages) ===
    for i in range(0, len(image_paths), 2):
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run()
        run1.add_picture(image_paths[i], width=Inches(5.6), height=Inches(4.2))
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(0)

        if i + 1 < len(image_paths):
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run()
            run2.add_picture(image_paths[i + 1], width=Inches(5.6), height=Inches(4.2))
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(0)

        if i + 2 < len(image_paths):
            doc.add_page_break()

    # Save to BytesIO for download
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# === Streamlit Inputs ===
weather = st.text_input("Weather Conditions (°C)")
max_dim = st.number_input("Max image dimension (px)", min_value=500, max_value=5000, value=1000)
quality = st.slider("JPEG Quality (1 = small file, 100 = high quality)", 1, 100, 90)
subcontractors = st.text_area("Subcontractors (one per line)").split("\n")
areas = st.text_area("Areas of Work (one per line)").split("\n")

# === Manage uploaded files using session state ===
if "uploaded_files" not in st.session_state:
    st.session_state.uploaded_files = []

uploaded = st.file_uploader("Upload Site Images", type=['jpg', 'jpeg'], accept_multiple_files=True)

# Prevent repopulation after clearing
if uploaded is not None and st.session_state.get("clear_triggered") is not True:
    st.session_state.uploaded_files = uploaded

# === Clear uploaded images button ===
if st.button("🗑️ Clear Uploaded Images"):
    st.session_state.uploaded_files = []
    st.session_state.clear_triggered = True
    st.rerun()

# Reset clear flag
if "clear_triggered" in st.session_state:
    del st.session_state.clear_triggered

# === Generate Report Button ===
if st.button("Generate Report"):
    if not st.session_state.uploaded_files:
        st.warning("Please upload at least one image.")
    else:
        with st.spinner("Generating report..."):
            uploaded_files = st.session_state.uploaded_files
            uploaded_files.sort(key=lambda f: extract_number(f.name))

            # Resize and save all images to temp paths
            temp_image_paths = []
            for file in uploaded_files:
                img = resize_image(file.read(), max_dim)
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
                img.save(temp_file.name, format="JPEG", quality=quality, optimize=True)
                temp_image_paths.append(temp_file.name)

            report_bytes = generate_report(temp_image_paths, weather, subcontractors, areas)
            st.success("Report generated!")
            st.download_button("📄 Download Report", report_bytes, file_name="Progress_Report.docx")
